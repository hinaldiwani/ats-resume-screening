"""app/utils/docx_parser.py"""
import docx


def extract_text_from_docx(file_path: str) -> str:
    document = docx.Document(file_path)
    text_parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text_parts.append(cell.text)
    return "\n".join(text_parts)
